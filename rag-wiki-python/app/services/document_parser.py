"""
增强型文档解析引擎
支持多格式文档解析、版面分析、表格提取、OCR识别
"""
import io
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from enum import Enum
from typing import List, Dict, Any, Optional, Tuple
import json

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """文档类型"""
    NATIVE_PDF = "native_pdf"
    SCANNED_PDF = "scanned_pdf"
    WORD = "word"
    EXCEL = "excel"
    PPT = "ppt"
    HTML = "html"
    TXT = "txt"
    UNKNOWN = "unknown"


class ParseStatus(Enum):
    """解析状态"""
    PENDING = "pending"
    PARSING = "parsing"
    COMPLETED = "completed"
    FAILED = "failed"
    NEED_REVIEW = "need_review"


@dataclass
class Table:
    """表格数据"""
    table_id: str
    page_num: int
    rows: int
    cols: int
    headers: List[str]
    data: List[List[str]]
    caption: Optional[str] = None


@dataclass
class Image:
    """图片数据"""
    image_id: str
    page_num: int
    description: Optional[str] = None
    ocr_text: Optional[str] = None


@dataclass
class TextBlock:
    """文本块"""
    block_id: str
    page_num: int
    text: str
    bbox: Tuple[float, float, float, float]  # x1, y1, x2, y2
    block_type: str  # title, paragraph, header, footer, etc.


@dataclass
class ParsedDocument:
    """解析后的文档"""
    document_id: str
    file_name: str
    file_type: DocumentType
    total_pages: int
    title: Optional[str] = None
    text_blocks: List[TextBlock] = None
    tables: List[Table] = None
    images: List[Image] = None
    metadata: Dict[str, Any] = None
    quality_score: float = 0.0
    status: ParseStatus = ParseStatus.PENDING
    error_message: Optional[str] = None
    
    def __post_init__(self):
        if self.text_blocks is None:
            self.text_blocks = []
        if self.tables is None:
            self.tables = []
        if self.images is None:
            self.images = []
        if self.metadata is None:
            self.metadata = {}


class BaseParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    def parse(self, file_content: bytes, file_name: str) -> ParsedDocument:
        """解析文档"""
        pass
    
    @abstractmethod
    def supports(self, file_type: DocumentType) -> bool:
        """是否支持该文档类型"""
        pass


class PDFParser(BaseParser):
    """PDF解析器"""
    
    def supports(self, file_type: DocumentType) -> bool:
        return file_type in [DocumentType.NATIVE_PDF, DocumentType.SCANNED_PDF]
    
    def parse(self, file_content: bytes, file_name: str) -> ParsedDocument:
        """解析PDF文档"""
        try:
            # 尝试使用PyMuPDF (fitz)
            import fitz
            
            doc = fitz.open(stream=file_content, filetype="pdf")
            parsed = ParsedDocument(
                document_id="",
                file_name=file_name,
                file_type=DocumentType.NATIVE_PDF,
                total_pages=len(doc)
            )
            
            # 提取文本块
            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("blocks")
                
                for block_idx, block in enumerate(blocks):
                    x1, y1, x2, y2, text, block_no, block_type = block
                    if text.strip():
                        text_block = TextBlock(
                            block_id=f"p{page_num}_b{block_idx}",
                            page_num=page_num + 1,
                            text=text.strip(),
                            bbox=(x1, y1, x2, y2),
                            block_type="paragraph" if block_type == 0 else "title"
                        )
                        parsed.text_blocks.append(text_block)
                
                # 提取表格（简化实现）
                tables = page.find_tables()
                for table_idx, table in enumerate(tables):
                    try:
                        df = table.to_pandas()
                        table_data = Table(
                            table_id=f"p{page_num}_t{table_idx}",
                            page_num=page_num + 1,
                            rows=len(df),
                            cols=len(df.columns),
                            headers=df.columns.tolist(),
                            data=df.values.tolist()
                        )
                        parsed.tables.append(table_data)
                    except Exception as e:
                        logger.warning(f"表格提取失败: {e}")
            
            doc.close()
            
            # 质量评估
            parsed.quality_score = self._evaluate_quality(parsed)
            parsed.status = ParseStatus.COMPLETED if parsed.quality_score > 0.6 else ParseStatus.NEED_REVIEW
            
            return parsed
            
        except ImportError:
            logger.error("PyMuPDF未安装，无法解析PDF")
            return self._create_error_doc(file_name, "PyMuPDF未安装")
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            return self._create_error_doc(file_name, str(e))
    
    def _evaluate_quality(self, doc: ParsedDocument) -> float:
        """评估解析质量"""
        if not doc.text_blocks:
            return 0.0
        
        # 基于文本提取率、结构完整性等评估
        total_text = sum(len(b.text) for b in doc.text_blocks)
        avg_block_length = total_text / len(doc.text_blocks) if doc.text_blocks else 0
        
        # 简单评分逻辑
        score = min(1.0, avg_block_length / 100)  # 平均每块100字符为满分
        return round(score, 2)
    
    def _create_error_doc(self, file_name: str, error: str) -> ParsedDocument:
        return ParsedDocument(
            document_id="",
            file_name=file_name,
            file_type=DocumentType.NATIVE_PDF,
            total_pages=0,
            status=ParseStatus.FAILED,
            error_message=error
        )


class WordParser(BaseParser):
    """Word文档解析器"""
    
    def supports(self, file_type: DocumentType) -> bool:
        return file_type == DocumentType.WORD
    
    def parse(self, file_content: bytes, file_name: str) -> ParsedDocument:
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_content))
            parsed = ParsedDocument(
                document_id="",
                file_name=file_name,
                file_type=DocumentType.WORD,
                total_pages=1  # docx不直接提供页数
            )
            
            # 提取段落
            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_block = TextBlock(
                        block_id=f"p{idx}",
                        page_num=1,
                        text=para.text.strip(),
                        bbox=(0, 0, 0, 0),
                        block_type="title" if para.style.name.startswith('Heading') else "paragraph"
                    )
                    parsed.text_blocks.append(text_block)
            
            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                try:
                    headers = [cell.text for cell in table.rows[0].cells] if table.rows else []
                    data = []
                    for row in table.rows[1:]:
                        row_data = [cell.text for cell in row.cells]
                        data.append(row_data)
                    
                    table_obj = Table(
                        table_id=f"t{table_idx}",
                        page_num=1,
                        rows=len(table.rows),
                        cols=len(table.rows[0].cells) if table.rows else 0,
                        headers=headers,
                        data=data
                    )
                    parsed.tables.append(table_obj)
                except Exception as e:
                    logger.warning(f"Word表格提取失败: {e}")
            
            parsed.quality_score = 0.85  # Word解析质量通常较高
            parsed.status = ParseStatus.COMPLETED
            return parsed
            
        except ImportError:
            return self._create_error_doc(file_name, "python-docx未安装")
        except Exception as e:
            logger.error(f"Word解析失败: {e}")
            return self._create_error_doc(file_name, str(e))
    
    def _create_error_doc(self, file_name: str, error: str) -> ParsedDocument:
        return ParsedDocument(
            document_id="",
            file_name=file_name,
            file_type=DocumentType.WORD,
            total_pages=0,
            status=ParseStatus.FAILED,
            error_message=error
        )


class ExcelParser(BaseParser):
    """Excel解析器"""
    
    def supports(self, file_type: DocumentType) -> bool:
        return file_type == DocumentType.EXCEL
    
    def parse(self, file_content: bytes, file_name: str) -> ParsedDocument:
        """解析Excel文档"""
        try:
            import pandas as pd
            
            # 读取所有sheet
            xls = pd.ExcelFile(io.BytesIO(file_content))
            parsed = ParsedDocument(
                document_id="",
                file_name=file_name,
                file_type=DocumentType.EXCEL,
                total_pages=len(xls.sheet_names)
            )
            
            for sheet_idx, sheet_name in enumerate(xls.sheet_names):
                df = pd.read_excel(xls, sheet_name=sheet_name)
                
                # 将DataFrame转为文本块
                text_content = f"Sheet: {sheet_name}\n"
                text_content += df.to_string(index=False)
                
                text_block = TextBlock(
                    block_id=f"s{sheet_idx}",
                    page_num=sheet_idx + 1,
                    text=text_content,
                    bbox=(0, 0, 0, 0),
                    block_type="table"
                )
                parsed.text_blocks.append(text_block)
                
                # 同时保存为表格对象
                table = Table(
                    table_id=f"s{sheet_idx}",
                    page_num=sheet_idx + 1,
                    rows=len(df),
                    cols=len(df.columns),
                    headers=df.columns.tolist(),
                    data=df.fillna('').values.tolist(),
                    caption=sheet_name
                )
                parsed.tables.append(table)
            
            parsed.quality_score = 0.9
            parsed.status = ParseStatus.COMPLETED
            return parsed
            
        except ImportError:
            return self._create_error_doc(file_name, "pandas未安装")
        except Exception as e:
            logger.error(f"Excel解析失败: {e}")
            return self._create_error_doc(file_name, str(e))
    
    def _create_error_doc(self, file_name: str, error: str) -> ParsedDocument:
        return ParsedDocument(
            document_id="",
            file_name=file_name,
            file_type=DocumentType.EXCEL,
            total_pages=0,
            status=ParseStatus.FAILED,
            error_message=error
        )


class DocumentParserService:
    """文档解析服务"""
    
    def __init__(self):
        self.parsers: List[BaseParser] = [
            PDFParser(),
            WordParser(),
            ExcelParser(),
        ]
    
    def detect_file_type(self, file_name: str) -> DocumentType:
        """检测文件类型"""
        ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
        
        type_mapping = {
            'pdf': DocumentType.NATIVE_PDF,
            'docx': DocumentType.WORD,
            'doc': DocumentType.WORD,
            'xlsx': DocumentType.EXCEL,
            'xls': DocumentType.EXCEL,
            'pptx': DocumentType.PPT,
            'ppt': DocumentType.PPT,
            'html': DocumentType.HTML,
            'htm': DocumentType.HTML,
            'txt': DocumentType.TXT,
        }
        
        return type_mapping.get(ext, DocumentType.UNKNOWN)
    
    def parse(self, file_content: bytes, file_name: str) -> ParsedDocument:
        """解析文档"""
        file_type = self.detect_file_type(file_name)
        logger.info(f"开始解析文档: {file_name}, 类型: {file_type.value}")
        
        if file_type == DocumentType.UNKNOWN:
            return ParsedDocument(
                document_id="",
                file_name=file_name,
                file_type=file_type,
                total_pages=0,
                status=ParseStatus.FAILED,
                error_message="不支持的文件类型"
            )
        
        # 查找合适的解析器
        for parser in self.parsers:
            if parser.supports(file_type):
                result = parser.parse(file_content, file_name)
                logger.info(f"文档解析完成: {file_name}, 质量评分: {result.quality_score}")
                return result
        
        return ParsedDocument(
            document_id="",
            file_name=file_name,
            file_type=file_type,
            total_pages=0,
            status=ParseStatus.FAILED,
            error_message="无可用解析器"
        )
    
    def chunk_document(self, parsed_doc: ParsedDocument, chunk_size: int = 500, 
                       chunk_overlap: int = 50) -> List[Dict[str, Any]]:
        """将解析后的文档分块"""
        chunks = []
        chunk_id = 0
        
        # 文本分块
        current_chunk = []
        current_length = 0
        
        for block in parsed_doc.text_blocks:
            text = block.text
            
            if current_length + len(text) > chunk_size and current_chunk:
                # 保存当前块
                chunk_text = "\n".join(current_chunk)
                chunks.append({
                    "chunk_id": f"chunk_{chunk_id}",
                    "text": chunk_text,
                    "page_num": block.page_num,
                    "type": "text"
                })
                chunk_id += 1
                
                # 保留重叠部分
                overlap_text = []
                overlap_length = 0
                for t in reversed(current_chunk):
                    if overlap_length + len(t) <= chunk_overlap:
                        overlap_text.insert(0, t)
                        overlap_length += len(t)
                    else:
                        break
                current_chunk = overlap_text
                current_length = overlap_length
            
            current_chunk.append(text)
            current_length += len(text)
        
        # 保存最后一个块
        if current_chunk:
            chunks.append({
                "chunk_id": f"chunk_{chunk_id}",
                "text": "\n".join(current_chunk),
                "page_num": parsed_doc.text_blocks[-1].page_num if parsed_doc.text_blocks else 1,
                "type": "text"
            })
        
        # 表格单独作为块
        for table in parsed_doc.tables:
            table_text = f"表格: {table.caption or ''}\n"
            table_text += " | ".join(table.headers) + "\n"
            for row in table.data[:10]:  # 限制行数
                table_text += " | ".join(str(cell) for cell in row) + "\n"
            
            chunks.append({
                "chunk_id": f"table_{table.table_id}",
                "text": table_text,
                "page_num": table.page_num,
                "type": "table"
            })
        
        return chunks


# 全局解析服务实例
parser_service = DocumentParserService()